# operation.py
from django.contrib import messages
from django.db import connection, DatabaseError, transaction
from django.http import JsonResponse, HttpResponseBadRequest
from django.shortcuts import redirect, render
from django.conf import settings
import logging
from supabase import create_client
import os
import io
import base64
import time

logger = logging.getLogger(__name__)
SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_SERVICE_KEY = os.getenv("SUPABASE_SERVICE_KEY")  # ⚠️ service role key, never expose to frontend
supabase = create_client(SUPABASE_URL, SUPABASE_SERVICE_KEY)
# =========================
# LOGIN via Postgres function (UPDATED to match your schema)
# =========================
def login_user(request):
    if request.method != "POST":
        return redirect("login")

    username = (request.POST.get("username") or "").strip()
    password = request.POST.get("password") or ""

    if not username or not password:
        messages.error(request, "Please enter username and password.")
        return redirect("login")

    try:
        with connection.cursor() as cur:
            cur.execute(
                """
                SELECT id, username, role, first_name, last_name, region_id, penro_id, cenro_id
                FROM auth_login(%s, %s);
                """,
                [username, password],
            )
            row = cur.fetchone()
    except DatabaseError as e:
        # Show detailed DB error only when DEBUG=True, to make troubleshooting easy
        if settings.DEBUG:
            logger.exception("DB error during auth_login()")
            msg = getattr(e, "pgerror", None) or str(e)
            messages.error(request, f"DB error: {msg}")
        else:
            messages.error(request, "Login service is temporarily unavailable.")
        return redirect("login")

    if not row:
        messages.error(request, "Username or password is incorrect.")
        return redirect("login")

    # Unpack and set session (matches updated auth_login RETURNS)
    user_id, uname, role, first_name, last_name, region_id, penro_id, cenro_id = row
    request.session["user_id"]    = user_id
    request.session["username"]   = uname
    request.session["role"]       = (role or "").strip().lower()
    request.session["first_name"] = first_name
    request.session["last_name"]  = last_name
    request.session["region_id"]  = region_id
    request.session["penro_id"]   = penro_id
    request.session["cenro_id"]   = cenro_id

    # Route by role (stored as 'Super Admin','Admin','PENRO','CENRO','Evaluator')
    r = request.session["role"]  # e.g. "super admin","admin","penro","cenro","evaluator"
    if r == "super admin":
        return redirect("SA-dashboard")
    elif r == "admin":
        return redirect("Admin-dashboard")
    elif r == "penro":
        return redirect("PENRO-dashboard")
    elif r == "cenro":
        return redirect("CENRO-dashboard")
    else:
        messages.error(request, "Undefined role. Contact support.")
        request.session.flush()
        return redirect("login")

# =========================
# LOGOUT
# =========================
def logout_user(request):
    request.session.flush()
    messages.success(request, "You have been logged out.")
    return redirect("login")

# =========================
# Helper: Check if user is logged in and get current user role
# =========================
def get_current_user_info(request):
    """Returns tuple of (user_role, region_id, penro_id, cenro_id) or (None, None, None, None)"""
    if not request.session.get("user_id"):
        return None, None, None, None
    
    role = request.session.get("role", "").lower()
    region_id = request.session.get("region_id")
    penro_id = request.session.get("penro_id")
    cenro_id = request.session.get("cenro_id")
    
    return role, region_id, penro_id, cenro_id

# =========================
# Helper: Get allowed roles for current user
# =========================
def get_allowed_roles_for_user(current_role):
    """Returns list of roles that current user can create"""
    role_hierarchy = {
        "super admin": ["Super Admin", "Admin"],
        "admin": ["Admin", "PENRO"],
        "penro": ["PENRO", "CENRO"],
        "cenro": ["CENRO", "Evaluator"],
        "evaluator": []  # Evaluators cannot create users
    }
    return role_hierarchy.get(current_role, [])

# =========================
# Helper: Get available offices based on current user's role and assignments
# =========================
def get_available_offices_for_user(current_role, current_region_id, current_penro_id, current_cenro_id):
    """Returns dict with available regions, penros, and cenros based on current user's permissions"""
    regions = []
    penros = []
    cenros = []
    
    with connection.cursor() as cur:
        if current_role == "super admin":
            # Super Admin can see all regions
            cur.execute("SELECT id, name FROM regions ORDER BY name;")
            regions = [{"id": r[0], "name": r[1]} for r in cur.fetchall()]
            
        elif current_role == "admin":
            # Admin can only see their assigned region and its PENROs
            if current_region_id:
                cur.execute("SELECT id, name FROM regions WHERE id = %s;", [current_region_id])
                regions = [{"id": r[0], "name": r[1]} for r in cur.fetchall()]
                
                cur.execute("SELECT id, name FROM penros WHERE region_id = %s ORDER BY name;", [current_region_id])
                penros = [{"id": r[0], "name": r[1]} for r in cur.fetchall()]
                
        elif current_role == "penro":
            # PENRO can only see their assigned PENRO and its CENROs
            if current_penro_id:
                cur.execute("SELECT id, name FROM penros WHERE id = %s;", [current_penro_id])
                penros = [{"id": r[0], "name": r[1]} for r in cur.fetchall()]
                
                cur.execute("SELECT id, name FROM cenros WHERE penro_id = %s ORDER BY name;", [current_penro_id])
                cenros = [{"id": r[0], "name": r[1]} for r in cur.fetchall()]
                
        elif current_role == "cenro":
            # CENRO can only see their assigned CENRO
            if current_cenro_id:
                cur.execute("SELECT id, name FROM cenros WHERE id = %s;", [current_cenro_id])
                cenros = [{"id": r[0], "name": r[1]} for r in cur.fetchall()]
    
    return {
        "regions": regions,
        "penros": penros,
        "cenros": cenros
    }

# =========================
# Helper: Validate office assignment based on role and current user permissions
# =========================
def validate_office_assignment(target_role, region_id, penro_id, cenro_id, current_user):
    """Validates if the current user can assign the target role to the specified offices"""
    current_role, current_region_id, current_penro_id, current_cenro_id = current_user
    
    # Convert string IDs to integers or None
    def to_int_or_none(v):
        try:
            return int(v) if v not in (None, "", "null") else None
        except Exception:
            return None
    
    r = to_int_or_none(region_id)
    p = to_int_or_none(penro_id)
    c = to_int_or_none(cenro_id)
    
    # Check role-specific office requirements
    if target_role == "Super Admin":
        if r or p or c:
            return False, "Super Admin cannot be assigned to any office."
    elif target_role == "Admin":
        if not r or p or c:
            return False, "Admin must be assigned to exactly one Region."
        # Check if current user can assign to this region
        if current_role == "super admin":
            pass  # Super admin can assign to any region
        elif current_role == "admin" and current_region_id == r:
            pass  # Admin can assign to their own region
        else:
            return False, "You don't have permission to assign Admin to this region."
    elif target_role == "PENRO":
        if r or not p or c:
            return False, "PENRO must be assigned to exactly one PENRO office."
        # Check if current user can assign to this PENRO
        if current_role == "super admin":
            pass  # Super admin can assign to any PENRO
        elif current_role == "admin":
            # Check if this PENRO belongs to admin's region
            with connection.cursor() as cur:
                cur.execute("SELECT region_id FROM penros WHERE id = %s;", [p])
                row = cur.fetchone()
                if not row or row[0] != current_region_id:
                    return False, "You can only assign PENRO within your region."
        elif current_role == "penro" and current_penro_id == p:
            pass  # PENRO can assign to their own office
        else:
            return False, "You don't have permission to assign PENRO to this office."
    elif target_role == "CENRO":
        if r or p or not c:
            return False, "CENRO must be assigned to exactly one CENRO office."
        # Check if current user can assign to this CENRO
        if current_role == "super admin":
            pass  # Super admin can assign to any CENRO
        elif current_role == "admin":
            # Check if this CENRO belongs to admin's region
            with connection.cursor() as cur:
                cur.execute("""
                    SELECT p.region_id 
                    FROM cenros c 
                    JOIN penros p ON c.penro_id = p.id 
                    WHERE c.id = %s;
                """, [c])
                row = cur.fetchone()
                if not row or row[0] != current_region_id:
                    return False, "You can only assign CENRO within your region."
        elif current_role == "penro":
            # Check if this CENRO belongs to PENRO's office
            with connection.cursor() as cur:
                cur.execute("SELECT penro_id FROM cenros WHERE id = %s;", [c])
                row = cur.fetchone()
                if not row or row[0] != current_penro_id:
                    return False, "You can only assign CENRO within your PENRO."
        elif current_role == "cenro" and current_cenro_id == c:
            pass  # CENRO can assign to their own office
        else:
            return False, "You don't have permission to assign CENRO to this office."
    elif target_role == "Evaluator":
        if not any([r, p, c]) or sum([bool(r), bool(p), bool(c)]) != 1:
            return False, "Evaluator must be assigned to exactly one office (Region OR PENRO OR CENRO)."
        # Check permissions based on which office is assigned
        if c:  # Assigned to CENRO
            if current_role == "cenro" and current_cenro_id == c:
                pass  # CENRO can assign evaluator to their office
            else:
                return False, "You can only assign Evaluator to your own CENRO office."
        elif p:  # Assigned to PENRO
            if current_role == "penro" and current_penro_id == p:
                pass  # PENRO can assign evaluator to their office
            else:
                return False, "You can only assign Evaluator to your own PENRO office."
        elif r:  # Assigned to Region
            if current_role == "admin" and current_region_id == r:
                pass  # Admin can assign evaluator to their region
            else:
                return False, "You can only assign Evaluator to your own region."
    
    return True, None

# =========================
# Helpers: fetch options (SQL) - Updated for hierarchical access
# =========================
def _fetch_regions():
    with connection.cursor() as cur:
        cur.execute("SELECT id, name FROM regions ORDER BY name;")
        return [{"id": r[0], "name": r[1]} for r in cur.fetchall()]

# =========================
# CREATE ACCOUNT (Updated with hierarchical permissions and auto-assignment)
# =========================
def create_account(request):
    # Check if user is logged in
    current_role, current_region_id, current_penro_id, current_cenro_id = get_current_user_info(request)
    if not current_role:
        messages.error(request, "You must be logged in to create accounts.")
        return redirect("login")
    
    # Check if user has permission to create accounts
    allowed_roles = get_allowed_roles_for_user(current_role)
    if not allowed_roles:
        messages.error(request, "You don't have permission to create user accounts.")
        return redirect("login")  # or appropriate dashboard
    
    if request.method == "POST":
        # Required
        first_name   = (request.POST.get("first_name") or "").strip()
        last_name    = (request.POST.get("last_name") or "").strip()
        gender       = (request.POST.get("gender") or "").strip()
        email        = (request.POST.get("email") or "").strip()
        role         = (request.POST.get("role") or "").strip()
        username     = (request.POST.get("username") or "").strip()
        password     = request.POST.get("password") or ""
        # Optional
        phone_number = (request.POST.get("phone_number") or "").strip() or None
        profile_pic  = (request.POST.get("profile_pic") or "").strip() or None

        # Basic validation
        errors = []
        if not first_name: errors.append("First name is required.")
        if not last_name:  errors.append("Last name is required.")
        if gender not in ("Male", "Female", "Other"):
            errors.append("Please pick a valid gender.")
        if not email:     errors.append("Email is required.")
        if role not in allowed_roles:
            errors.append(f"You can only create users with these roles: {', '.join(allowed_roles)}")
        if not username:  errors.append("Username is required.")
        if not password:  errors.append("Password is required.")
        if errors:
            for e in errors: messages.error(request, e)
            return redirect("account-create")

        # Auto-assign office IDs based on current user's role and target role
        r = None  # region_id
        p = None  # penro_id  
        c = None  # cenro_id

        if role == "Super Admin":
            # Super Admin has no office assignments
            r = p = c = None
            
        elif current_role == "super admin":
            # Super Admin creating other roles - get from form
            if role == "Admin":
                # For Admin creation by Super Admin, require region selection from form
                region_id_from_form = request.POST.get("region_id")
                try:
                    r = int(region_id_from_form) if region_id_from_form not in (None, "", "null") else None
                except:
                    r = None
                
                if not r:
                    messages.error(request, "Please select a region for the Admin user.")
                    return redirect("account-create")
                
                p = c = None
            elif role == "PENRO":
                # Super Admin can also create PENRO - get penro_id from form
                penro_id_from_form = request.POST.get("penro_id")
                try:
                    p = int(penro_id_from_form) if penro_id_from_form not in (None, "", "null") else None
                except:
                    p = None
                
                if not p:
                    messages.error(request, "Please select a PENRO office.")
                    return redirect("account-create")
                
                r = c = None
            
        elif current_role == "admin":
            # Admin can create Admin or PENRO - both should inherit admin's region
            if role == "Admin":
                r = current_region_id
                p = c = None
            elif role == "PENRO":
                # For PENRO creation, get penro_id from form and inherit admin's region
                penro_id_from_form = request.POST.get("penro_id")
                try:
                    p = int(penro_id_from_form) if penro_id_from_form not in (None, "", "null") else None
                except:
                    p = None
                
                if not p:
                    messages.error(request, "Please select a PENRO office.")
                    return redirect("account-create")
                    
                # Verify this PENRO belongs to admin's region
                with connection.cursor() as cur:
                    cur.execute("SELECT region_id FROM penros WHERE id = %s;", [p])
                    row = cur.fetchone()
                    if not row or row[0] != current_region_id:
                        messages.error(request, "Selected PENRO does not belong to your region.")
                        return redirect("account-create")
                
                # PENRO should inherit admin's region AND have penro assignment
                r = current_region_id  # Inherit region from admin
                c = None
                
        elif current_role == "penro":
            # PENRO can create PENRO or CENRO - both should inherit penro's assignments
            if role == "PENRO":
                # Inherit both region and penro from current PENRO user
                r = current_region_id
                p = current_penro_id
                c = None
            elif role == "CENRO":
                # For CENRO creation, get cenro_id from form and inherit penro's region
                cenro_id_from_form = request.POST.get("cenro_id")
                try:
                    c = int(cenro_id_from_form) if cenro_id_from_form not in (None, "", "null") else None
                except:
                    c = None
                    
                if not c:
                    messages.error(request, "Please select a CENRO office.")
                    return redirect("account-create")
                    
                # Verify this CENRO belongs to penro's office
                with connection.cursor() as cur:
                    cur.execute("SELECT penro_id FROM cenros WHERE id = %s;", [c])
                    row = cur.fetchone()
                    if not row or row[0] != current_penro_id:
                        messages.error(request, "Selected CENRO does not belong to your PENRO.")
                        return redirect("account-create")
                
                # Get the region_id for this CENRO through PENRO
                with connection.cursor() as cur:
                    cur.execute("""
                        SELECT p.region_id 
                        FROM cenros c 
                        JOIN penros p ON c.penro_id = p.id 
                        WHERE c.id = %s;
                    """, [c])
                    row = cur.fetchone()
                    if row:
                        r = row[0]  # Inherit region through PENRO
                
                # CENRO should inherit region, penro, AND have cenro assignment
                p = current_penro_id  # Inherit penro from current user
                
        elif current_role == "cenro":
            # CENRO can create CENRO or Evaluator - both should inherit cenro's full hierarchy
            if role == "CENRO":
                # Inherit full hierarchy: region, penro, and cenro
                r = current_region_id
                p = current_penro_id  
                c = current_cenro_id
            elif role == "Evaluator":
                # Evaluator assigned to current CENRO inherits full hierarchy
                r = current_region_id
                p = current_penro_id
                c = current_cenro_id

        try:
            with transaction.atomic():
                with connection.cursor() as cur:
                    cur.execute(
                        """
                                                SELECT create_user(
                                %s, %s, %s, %s,   -- first_name, last_name, gender, email
                                %s,               -- phone_number
                                %s,               -- role
                                %s, %s,           -- username, password
                                %s,               -- profile_pic
                                %s, %s, %s,       -- region_id, penro_id, cenro_id
                                %s                -- p_hash_password
                            );
                        """,
                        [
                            first_name, last_name, gender, email,
                            phone_number, role, username,
                            password, profile_pic,
                            r, p, c,
                            True,
                        ],
                    )
                    new_id = cur.fetchone()[0]

            messages.success(request, f"Account successfully created! User can now log in.")
            return redirect("account-create")  # Stay on page to create more users

        except DatabaseError as e:
            msg = getattr(e, "pgerror", None) or str(e)
            if "exists" in msg.lower():
                messages.error(request, "Email or username already exists.")
            else:
                messages.error(request, f"DB error: {msg}")
            return redirect("account-create")

    # GET → show form with appropriate options based on current user's role
    available_offices = get_available_offices_for_user(current_role, current_region_id, current_penro_id, current_cenro_id)
    
    # Get current user's office names for display
    current_user_region_name = None
    current_user_penro_name = None
    current_user_cenro_name = None
    
    with connection.cursor() as cur:
        if current_region_id:
            cur.execute("SELECT name FROM regions WHERE id = %s;", [current_region_id])
            row = cur.fetchone()
            if row:
                current_user_region_name = row[0]
        
        if current_penro_id:
            cur.execute("SELECT name FROM penros WHERE id = %s;", [current_penro_id])
            row = cur.fetchone()
            if row:
                current_user_penro_name = row[0]
        
        if current_cenro_id:
            cur.execute("SELECT name FROM cenros WHERE id = %s;", [current_cenro_id])
            row = cur.fetchone()
            if row:
                current_user_cenro_name = row[0]
    
    context = {
        "regions": available_offices["regions"],
        "penros": available_offices["penros"],
        "cenros": available_offices["cenros"],
        "roles": allowed_roles,
        "genders": ["Male", "Female", "Other"],
        "current_user_role": current_role.title(),
        # Current user's office assignments for auto-inheritance
        "current_user_region_id": current_region_id,
        "current_user_penro_id": current_penro_id,
        "current_user_cenro_id": current_cenro_id,
        "current_user_region_name": current_user_region_name,
        "current_user_penro_name": current_user_penro_name,
        "current_user_cenro_name": current_user_cenro_name,
    }
    return render(request, "create_account.html", context)

# =========================
# AJAX APIs for cascading selects (Updated with permission checks)
# =========================
def api_penros_by_region(request, region_id):
    current_role, current_region_id, current_penro_id, current_cenro_id = get_current_user_info(request)
    if not current_role:
        return HttpResponseBadRequest("Not authenticated")
    
    try:
        rid = int(region_id)
    except Exception:
        return HttpResponseBadRequest("Invalid region id")
    
    # Check if user has access to this region
    if current_role == "admin" and current_region_id != rid:
        return HttpResponseBadRequest("Access denied to this region")
    
    with connection.cursor() as cur:
        if current_role in ["super admin", "admin"]:
            cur.execute("SELECT id, name FROM penros WHERE region_id=%s ORDER BY name;", [rid])
        else:
            # For other roles, return empty list
            return JsonResponse({"items": []})
        
        items = [{"id": r[0], "name": r[1]} for r in cur.fetchall()]
    return JsonResponse({"items": items})

def api_cenros_by_penro(request, penro_id):
    current_role, current_region_id, current_penro_id, current_cenro_id = get_current_user_info(request)
    if not current_role:
        return HttpResponseBadRequest("Not authenticated")
    
    try:
        pid = int(penro_id)
    except Exception:
        return HttpResponseBadRequest("Invalid penro id")
    
    # Check if user has access to this PENRO
    if current_role == "admin":
        # Check if this PENRO belongs to admin's region
        with connection.cursor() as cur:
            cur.execute("SELECT region_id FROM penros WHERE id = %s;", [pid])
            row = cur.fetchone()
            if not row or row[0] != current_region_id:
                return HttpResponseBadRequest("Access denied to this PENRO")
    elif current_role == "penro" and current_penro_id != pid:
        return HttpResponseBadRequest("Access denied to this PENRO")
    
    with connection.cursor() as cur:
        if current_role in ["super admin", "admin", "penro"]:
            cur.execute("SELECT id, name FROM cenros WHERE penro_id=%s ORDER BY name;", [pid])
        else:
            # For other roles, return empty list
            return JsonResponse({"items": []})
        
        items = [{"id": r[0], "name": r[1]} for r in cur.fetchall()]
    return JsonResponse({"items": items})
# =========================
# GET ENUMERATOR REPORTS for CENRO (WITH ALL FILTERS)
# =========================
def get_enumerator_reports(cenro_id=None, from_date=None, to_date=None, establishment_type=None, pa_id=None, establishment_status=None):
    """
    Fetch enumerator reports for a CENRO office with all filters.
    
    Args:
        cenro_id: Filter reports by CENRO office (None for all)
        from_date: Start date (datetime.date object or None)
        to_date: End date (datetime.date object or None)
        establishment_type: Filter by establishment type (None for all)
        pa_id: Filter by protected area ID (None for all)
        establishment_status: Filter by establishment status (None for all)
    
    Returns:
        List of report dictionaries
    """
    reports = []
    
    try:
        with connection.cursor() as cur:
            query = """
                SELECT 
                    er.id,
                    er.establishment_name,
                    er.proponent_name,
                    er.pa_name,
                    er.enumerator_name,
                    er.report_date,
                    er.informant_name,
                    er.remarks,
                    er.created_at,
                    u.first_name || ' ' || u.last_name as enumerator_full_name,
                    u.cenro_id,
                    ep.establishment_type,
                    er.pa_id,
                    ep.establishment_status
                FROM enumerators_report er
                LEFT JOIN users u ON er.enumerator_id = u.id
                LEFT JOIN establishment_profile ep ON er.establishment_id = ep.id
                WHERE 1=1
            """
            
            params = []
            
            if cenro_id:
                query += " AND u.cenro_id = %s"
                params.append(cenro_id)
            
            if from_date:
                query += " AND er.report_date >= %s"
                params.append(from_date)
            
            if to_date:
                query += " AND er.report_date <= %s"
                params.append(to_date)
            
            if establishment_type:
                query += " AND LOWER(TRIM(ep.establishment_type)) = LOWER(TRIM(%s))"
                params.append(establishment_type)
            
            if pa_id:
                query += " AND er.pa_id = %s"
                params.append(pa_id)
            
            if establishment_status:
                query += " AND LOWER(TRIM(ep.establishment_status)) = LOWER(TRIM(%s))"
                params.append(establishment_status)
            
            query += " ORDER BY er.report_date DESC, er.created_at DESC;"
            
            cur.execute(query, params)
            
            columns = [desc[0] for desc in cur.description]
            for row in cur.fetchall():
                reports.append(dict(zip(columns, row)))
                
    except DatabaseError as e:
        logger.error(f"Error fetching enumerator reports: {e}")
        
    return reports


def get_establishment_types_for_cenro(cenro_id):
    """
    Get list of establishment types for a specific CENRO.
    
    Args:
        cenro_id: The CENRO office ID
    
    Returns:
        List of establishment type strings
    """
    establishment_types = []
    
    try:
        with connection.cursor() as cur:
            query = """
                SELECT DISTINCT ep.establishment_type
                FROM enumerators_report er
                LEFT JOIN users u ON er.enumerator_id = u.id
                LEFT JOIN establishment_profile ep ON er.establishment_id = ep.id
                WHERE u.cenro_id = %s
                AND ep.establishment_type IS NOT NULL
                ORDER BY ep.establishment_type;
            """
            
            cur.execute(query, [cenro_id])
            
            for row in cur.fetchall():
                if row[0]:
                    establishment_types.append(row[0])
                
    except DatabaseError as e:
        logger.error(f"Error fetching establishment types: {e}")
        
    return establishment_types


def get_protected_areas_for_cenro(cenro_id):
    """
    Get list of all protected areas from Supabase.
    
    Args:
        cenro_id: The CENRO office ID (not used, kept for compatibility)
    
    Returns:
        List of dictionaries with id and name
    """
    protected_areas = []
    
    try:
        # Fetch all PA details from Supabase
        result = supabase.table('protected_areas').select('id, name').order('name').execute()
        protected_areas = result.data if result.data else []
                
    except Exception as e:
        logger.error(f"Error fetching protected areas: {e}")
        
    return protected_areas


def get_report_details(report_id, cenro_id=None):
    """
    Get detailed information about a specific enumerator report.
    
    Args:
        report_id: The report ID
        cenro_id: Optional CENRO ID to verify access
    
    Returns:
        Dictionary with complete report details or None if not found
    """
    try:
        with connection.cursor() as cur:
            # Select all enumerators_report columns and relevant joined fields. Using er.* makes it easier
            # to return every report column even if NULL, satisfying the requirement to "get all the form".
            query = """
                SELECT
                    er.*, 
                    u.first_name || ' ' || u.last_name AS enumerator_full_name,
                    u.cenro_id,
                    ep.establishment_type,
                    ep.establishment_status,
                    ep.description,
                    
                    ep.lot_status,
                    ep.land_classification,
                    ep.title_no,
                    ep.lot_no,
                    ep.lot_owner,
                    ep.area_covered,
                    ep.pa_zone,
                    ep.within_easement,
                    ep.tax_declaration_no,
                    ep.mayor_permit_no,
                    ep.mayor_permit_issued,
                    ep.mayor_permit_exp,
                    ep.business_permit_no,
                    ep.business_permit_issued,
                    ep.business_permit_exp,
                    ep.building_permit_no,
                    ep.building_permit_issued,
                    ep.building_permit_exp,
                    ep.pamb_resolution_no,
                    ep.pamb_date_issued,
                    ep.sapa_no,
                    ep.sapa_date_issued,
                    ep.pacbrma_no,
                    ep.pacbrma_date_issued,
                    ep.ecc_no,
                    ep.ecc_date_issued,
                    ep.discharge_permit_no,
                    ep.discharge_date_issued,
                    ep.pto_no,
                    ep.pto_date_issued,
                    ep.other_emb,
                    gti.image AS geo_image_url,
                    gti.latitude AS geo_latitude,
                    gti.longitude AS geo_longitude,
                    gti.location AS geo_location,
                    gti.captured_at AS geo_captured_at,
                    an.attested_by_name,
                    an.attested_by_position,
                    an.attested_by_signature,
                    an.noted_by_name,
                    an.noted_by_position,
                    an.noted_by_signature
                FROM enumerators_report er
                LEFT JOIN users u ON er.enumerator_id = u.id
                LEFT JOIN establishment_profile ep ON er.establishment_id = ep.id
                LEFT JOIN geo_tagged_images gti ON er.geo_tagged_image_id = gti.id
                LEFT JOIN attestation_notations an ON er.attestation_id = an.id
                WHERE er.id = %s
            """

            params = [report_id]

            if cenro_id:
                query += " AND u.cenro_id = %s"
                params.append(cenro_id)

            cur.execute(query, params)
            row = cur.fetchone()

            if not row:
                return None

            # Build dictionary using cursor description so every column from er.* is present
            columns = [desc[0] for desc in cur.description]
            data = dict(zip(columns, row))

            # If proponent_name is missing, try plausible fallbacks in order:
            # 1) If report links to a leased property profile (profile_id), use its proponent_name
            # 2) If report has proponent_id, fetch from proponents table
            if (not data.get('proponent_name') or str(data.get('proponent_name')).strip() == ''):
                # Try leased property profile (common in this schema)
                profile_id = data.get('profile_id') or data.get('profile')
                if profile_id:
                    try:
                        cur.execute("SELECT proponent_name FROM leasedpropertyprofile WHERE id = %s;", [profile_id])
                        p_row = cur.fetchone()
                        if p_row and p_row[0]:
                            data['proponent_name'] = p_row[0]
                    except Exception:
                        # ignore lookup errors and continue to next fallback
                        pass

                # Fallback: try proponents table if still missing
                if (not data.get('proponent_name') or str(data.get('proponent_name')).strip() == '') and data.get('proponent_id'):
                    try:
                        cur.execute("SELECT name FROM proponents WHERE id = %s;", [data.get('proponent_id')])
                        p_row = cur.fetchone()
                        if p_row and p_row[0]:
                            data['proponent_name'] = p_row[0]
                    except Exception:
                        # ignore lookup errors - leave proponent_name as-is
                        pass

            # Build permits array using keys that may be present from ep.*
            permits = []
            if data.get('mayor_permit_no') or data.get('mayor_permit_issued') or data.get('mayor_permit_exp'):
                permits.append({
                    'name': "Mayor's Permit",
                    'number': data.get('mayor_permit_no'),
                    'issued': data.get('mayor_permit_issued').isoformat() if data.get('mayor_permit_issued') else None,
                    'expiry': data.get('mayor_permit_exp').isoformat() if data.get('mayor_permit_exp') else None
                })

            if data.get('business_permit_no') or data.get('business_permit_issued') or data.get('business_permit_exp'):
                permits.append({
                    'name': 'Business Permit',
                    'number': data.get('business_permit_no'),
                    'issued': data.get('business_permit_issued').isoformat() if data.get('business_permit_issued') else None,
                    'expiry': data.get('business_permit_exp').isoformat() if data.get('business_permit_exp') else None
                })

            if data.get('building_permit_no') or data.get('building_permit_issued') or data.get('building_permit_exp'):
                permits.append({
                    'name': 'Building Permit',
                    'number': data.get('building_permit_no'),
                    'issued': data.get('building_permit_issued').isoformat() if data.get('building_permit_issued') else None,
                    'expiry': data.get('building_permit_exp').isoformat() if data.get('building_permit_exp') else None
                })

            if data.get('pamb_resolution_no') or data.get('pamb_date_issued'):
                permits.append({
                    'name': 'PAMB Resolution',
                    'number': data.get('pamb_resolution_no'),
                    'issued': data.get('pamb_date_issued').isoformat() if data.get('pamb_date_issued') else None,
                    'expiry': None
                })

            if data.get('sapa_no') or data.get('sapa_date_issued'):
                permits.append({
                    'name': 'SAPA',
                    'number': data.get('sapa_no'),
                    'issued': data.get('sapa_date_issued').isoformat() if data.get('sapa_date_issued') else None,
                    'expiry': None
                })

            if data.get('pacbrma_no') or data.get('pacbrma_date_issued'):
                permits.append({
                    'name': 'PACBRMA',
                    'number': data.get('pacbrma_no'),
                    'issued': data.get('pacbrma_date_issued').isoformat() if data.get('pacbrma_date_issued') else None,
                    'expiry': None
                })

            if data.get('ecc_no') or data.get('ecc_date_issued'):
                permits.append({
                    'name': 'Environmental Compliance Certificate (ECC)',
                    'number': data.get('ecc_no'),
                    'issued': data.get('ecc_date_issued').isoformat() if data.get('ecc_date_issued') else None,
                    'expiry': None
                })

            if data.get('discharge_permit_no') or data.get('discharge_date_issued'):
                permits.append({
                    'name': 'Discharge Permit',
                    'number': data.get('discharge_permit_no'),
                    'issued': data.get('discharge_date_issued').isoformat() if data.get('discharge_date_issued') else None,
                    'expiry': None
                })

            if data.get('pto_no') or data.get('pto_date_issued'):
                permits.append({
                    'name': 'Permit to Operate (PTO)',
                    'number': data.get('pto_no'),
                    'issued': data.get('pto_date_issued').isoformat() if data.get('pto_date_issued') else None,
                    'expiry': None
                })

            # Helper to build full URL for signatures stored as relative paths
            def build_signature_url(sig_path):
                if not sig_path:
                    return None
                # If already absolute URL, return as-is
                if sig_path.startswith('http://') or sig_path.startswith('https://') or sig_path.startswith('data:'):
                    return sig_path
                # Build Supabase public URL
                bucket = os.getenv('SUPABASE_BUCKET', 'geo-tagged-photos')
                return f"{SUPABASE_URL}/storage/v1/object/public/{bucket}/{sig_path}"

            # Prepare response ensuring key names expected by frontend are present
            report_details = {
                # fields coming from er.* - ensure presence even if None
                'id': data.get('id'),
                'establishment_id': data.get('establishment_id'),
                'establishment_name': data.get('establishment_name'),
                'proponent_id': data.get('proponent_id') or data.get('proponent_id'),
                'proponent_name': data.get('proponent_name'),
                'pa_id': data.get('pa_id'),
                'pa_name': data.get('pa_name'),
                'enumerator_id': data.get('enumerator_id'),
                'enumerator_name': data.get('enumerator_name') or data.get('enumerator_full_name'),
                'geo_tagged_image_id': data.get('geo_tagged_image_id'),
                'report_date': data.get('report_date').isoformat() if data.get('report_date') else None,
                'enumerator_signature_date': data.get('enumerator_signature_date').isoformat() if data.get('enumerator_signature_date') else None,
                'informant_signature_date': data.get('informant_signature_date').isoformat() if data.get('informant_signature_date') else None,
                'enumerator_signature': build_signature_url(data.get('enumerator_signature')),
                'informant_signature': build_signature_url(data.get('informant_signature')),
                'informant_name': data.get('informant_name'),
                'remarks': data.get('remarks'),
                'created_at': data.get('created_at').isoformat() if data.get('created_at') else None,
                'updated_at': data.get('updated_at').isoformat() if data.get('updated_at') else None,

                # establishment_profile derived fields
                'establishment_type': data.get('establishment_type'),
                'establishment_status': data.get('establishment_status'),
                'description': data.get('description'),
                'lot_status': data.get('lot_status'),
                'land_classification': data.get('land_classification'),
                'title_no': data.get('title_no'),
                'lot_no': data.get('lot_no'),
                'lot_owner': data.get('lot_owner'),
                'area_covered': data.get('area_covered'),
                'pa_zone': data.get('pa_zone'),
                'within_easement': data.get('within_easement'),
                'tax_declaration_no': data.get('tax_declaration_no'),

                'permits': permits,
                'other_emb': data.get('other_emb'),

                # geo image fields
                'geo_image_url': data.get('geo_image_url'),
                'latitude': data.get('geo_latitude') or data.get('latitude'),
                'longitude': data.get('geo_longitude') or data.get('longitude'),
                'location': data.get('geo_location') or data.get('location'),
                'geo_captured_at': data.get('geo_captured_at').isoformat() if data.get('geo_captured_at') else None,

                # attestation fields - build full URLs
                'attestation_id': data.get('attestation_id'),
                'attested_by_name': data.get('attested_by_name'),
                'attested_by_position': data.get('attested_by_position'),
                'attested_by_signature': build_signature_url(data.get('attested_by_signature')),
                'noted_by_name': data.get('noted_by_name'),
                'noted_by_position': data.get('noted_by_position'),
                'noted_by_signature': build_signature_url(data.get('noted_by_signature'))
            }

            return report_details
                
    except DatabaseError as e:
        logger.error(f"Error fetching report details: {e}")
        return None
    

def get_report_images(report_id):
    """
    Fetch images related to a report from Supabase storage/join table.

    Returns a list of dicts with keys: id, image, latitude, longitude, location,
    captured_at, qr_code, is_primary, image_sequence
    """
    images = []
    if not report_id:
        return images

    try:
        resp = supabase.table("reported_images") \
            .select("image_id,is_primary,image_sequence,geo_tagged_images(id,image,latitude,longitude,location,captured_at,qr_code)") \
            .eq("report_id", report_id) \
            .eq("report_type", "enumerator") \
            .execute()

        reported = getattr(resp, 'data', None) or []

        # Some Supabase client versions return unsorted results; sort by image_sequence if available
        try:
            reported.sort(key=lambda x: (x.get('image_sequence') is None, x.get('image_sequence') or 0))
        except Exception:
            pass

        for ri in reported:
            geo = ri.get('geo_tagged_images')
            if not geo:
                continue
            images.append({
                'id': geo.get('id'),
                'image': geo.get('image'),
                'latitude': geo.get('latitude'),
                'longitude': geo.get('longitude'),
                'location': geo.get('location'),
                'captured_at': geo.get('captured_at'),
                'qr_code': geo.get('qr_code'),
                'is_primary': ri.get('is_primary', False),
                'image_sequence': ri.get('image_sequence', None),
            })

    except Exception as e:
        logger.exception(f"Error fetching report images for report_id=%s: %s", report_id, e)

    return images


def save_notation(report_id, noted_by_name, noted_by_position, signature_dataurl, current_user_id=None):
    """Save notation record and upload signature image to Supabase storage.

    Returns (True, public_url) on success, or (False, error_message) on failure.
    """
    try:
        if not report_id:
            return False, 'Invalid report id'

        if not signature_dataurl or not signature_dataurl.startswith('data:'):
            return False, 'Invalid signature data'

        header, encoded = signature_dataurl.split(',', 1)
        try:
            data = base64.b64decode(encoded)
        except Exception as e:
            return False, f'Decoding error: {e}'

        bucket = os.getenv('SUPABASE_BUCKET', 'geo-tagged-photos')
        filename = f"attestation/report_{report_id}_noted_{int(time.time())}.png"

        try:
            from_call = supabase.storage.from_(bucket)
            upload_resp = from_call.upload(filename, data, {"content-type": "image/png"})
            signature_url_to_store = filename
        except Exception as e:
            logger.exception('Supabase upload failed: %s', e)
            return False, f'Upload failed: {str(e)}'

        with connection.cursor() as cur:
            cur.execute("SELECT attestation_id FROM enumerators_report WHERE id = %s;", [report_id])
            row = cur.fetchone()
            existing_id = row[0] if row else None

            if existing_id:
                cur.execute(
                    """
                    UPDATE attestation_notations
                    SET noted_by_name = %s,
                        noted_by_position = %s,
                        noted_by_signature = %s
                    WHERE id = %s
                    RETURNING id;
                    """,
                    [noted_by_name, noted_by_position, signature_url_to_store, existing_id]
                )
                cur.fetchone()
            else:
                cur.execute(
                    """
                    INSERT INTO attestation_notations (noted_by_name, noted_by_position, noted_by_signature)
                    VALUES (%s, %s, %s)
                    RETURNING id;
                    """,
                    [noted_by_name, noted_by_position, signature_url_to_store]
                )
                new_id = cur.fetchone()[0]
                cur.execute("UPDATE enumerators_report SET attestation_id = %s WHERE id = %s;", [new_id, report_id])

        full_url = f"{SUPABASE_URL}/storage/v1/object/public/{bucket}/{signature_url_to_store}"
        return True, full_url

    except DatabaseError as e:
        logger.exception('DB error saving notation: %s', e)
        return False, str(e)
    except Exception as e:
        logger.exception('Unexpected error saving notation: %s', e)
        return False, str(e)


def save_attestation(report_id, attested_by_name, attested_by_position, signature_dataurl, current_user_id=None):
    """Save attestation record and upload signature image to Supabase storage.

    Returns (True, public_url) on success, or (False, error_message) on failure.
    """
    try:
        if not report_id:
            return False, 'Invalid report id'

        # Decode data URL
        if not signature_dataurl or not signature_dataurl.startswith('data:'):
            return False, 'Invalid signature data'

        header, encoded = signature_dataurl.split(',', 1)
        try:
            data = base64.b64decode(encoded)
        except Exception as e:
            return False, f'Decoding error: {e}'

        # Prepare file path in attestation folder
        bucket = os.getenv('SUPABASE_BUCKET', 'geo-tagged-photos')
        filename = f"attestation/report_{report_id}_attested_{int(time.time())}.png"

        # Upload to Supabase storage
        try:
            from_call = supabase.storage.from_(bucket)
            upload_resp = from_call.upload(filename, data, {"content-type": "image/png"})
            signature_url_to_store = filename
        except Exception as e:
            logger.exception('Supabase upload failed: %s', e)
            return False, f'Upload failed: {str(e)}'

        # Insert or update attestation_notations and link to enumerators_report
        with connection.cursor() as cur:
            # Check existing attestation_id
            cur.execute("SELECT attestation_id FROM enumerators_report WHERE id = %s;", [report_id])
            row = cur.fetchone()
            existing_id = row[0] if row else None

            if existing_id:
                cur.execute(
                    """
                    UPDATE attestation_notations
                    SET attested_by_name = %s,
                        attested_by_position = %s,
                        attested_by_signature = %s
                    WHERE id = %s
                    RETURNING id;
                    """,
                    [attested_by_name, attested_by_position, signature_url_to_store, existing_id]
                )
                cur.fetchone()
            else:
                cur.execute(
                    """
                    INSERT INTO attestation_notations (attested_by_name, attested_by_position, attested_by_signature)
                    VALUES (%s, %s, %s)
                    RETURNING id;
                    """,
                    [attested_by_name, attested_by_position, signature_url_to_store]
                )
                new_id = cur.fetchone()[0]
                cur.execute("UPDATE enumerators_report SET attestation_id = %s WHERE id = %s;", [new_id, report_id])

        # Return full URL for response
        full_url = f"{SUPABASE_URL}/storage/v1/object/public/{bucket}/{signature_url_to_store}"
        return True, full_url

    except DatabaseError as e:
        logger.exception('DB error saving attestation: %s', e)
        return False, str(e)
    except Exception as e:
        logger.exception('Unexpected error saving attestation: %s', e)
        return False, str(e)




def get_activity_logs():
    """Fetch activity logs via DB function `get_activity_logs()` using Django connection."""
    try:
        with connection.cursor() as cur:
            cur.execute("SELECT * FROM get_activity_logs();")
            logs = cur.fetchall()

            log_list = []
            for row in logs:
                log_list.append({
                    "task": row[0],
                    "user_id": row[1],
                    "name": row[2],
                    "timestamp": row[3],
                })

            return log_list
    except Exception as e:
        logger.exception("Error fetching activity logs: %s", e)
        return []


# =========================
# PROTECTED AREAS MANAGEMENT
# =========================
def add_protected_area(name, file_obj):
    """Add protected area with file upload to Supabase."""
    try:
        # Determine file type
        file_name = file_obj.name.lower()
        if file_name.endswith('.kml'):
            file_type = 'kml'
        elif file_name.endswith('.zip'):
            file_type = 'shp'
        else:
            return False, 'Invalid file type. Upload KML or ZIP (containing shapefile).'

        # Upload file to Supabase storage
        bucket = os.getenv('SUPABASE_BUCKET', 'geo-tagged-photos')
        file_path = f"protected-areas/{int(time.time())}_{file_obj.name}"
        
        try:
            file_data = file_obj.read()
            from_call = supabase.storage.from_(bucket)
            upload_resp = from_call.upload(file_path, file_data, {"content-type": file_obj.content_type or "application/octet-stream"})
        except Exception as e:
            logger.exception('Supabase file upload failed: %s', e)
            return False, f'File upload failed: {str(e)}'

        # Insert into Supabase table
        try:
            result = supabase.table('protected_areas').insert({
                'name': name,
                'file_type': file_type,
                'file_path': file_path
            }).execute()
            
            return True, 'Protected area added successfully'
        except Exception as e:
            logger.exception('Supabase insert failed: %s', e)
            return False, f'Database insert failed: {str(e)}'

    except Exception as e:
        logger.exception('Error adding protected area: %s', e)
        return False, str(e)


def get_protected_areas():
    """Fetch all protected areas from Supabase."""
    try:
        result = supabase.table('protected_areas').select('*').order('created_at', desc=True).execute()
        return result.data if result.data else []
    except Exception as e:
        logger.exception('Error fetching protected areas: %s', e)
        return []


def delete_protected_area(pa_id):
    """Delete protected area and its file from Supabase."""
    try:
        # Get file path before deleting
        result = supabase.table('protected_areas').select('file_path').eq('id', pa_id).execute()
        
        if result.data and len(result.data) > 0:
            file_path = result.data[0].get('file_path')
            
            # Delete file from storage
            if file_path:
                try:
                    bucket = os.getenv('SUPABASE_BUCKET', 'geo-tagged-photos')
                    supabase.storage.from_(bucket).remove([file_path])
                except Exception as e:
                    logger.warning('Failed to delete file from storage: %s', e)
            
            # Delete from table
            supabase.table('protected_areas').delete().eq('id', pa_id).execute()
            return True, 'Protected area deleted successfully'
        else:
            return False, 'Protected area not found'
            
    except Exception as e:
        logger.exception('Error deleting protected area: %s', e)
        return False, str(e)


def export_reports(reports, format_type):
    """Export reports to PDF, Word, or Excel format"""
    from io import BytesIO
    from django.http import HttpResponse
    
    if format_type == 'excel':
        from openpyxl import Workbook
        wb = Workbook()
        ws = wb.active
        ws.title = 'Reports'
        ws.append(['Report ID', 'Date', 'Establishment', 'Type', 'Status', 'Protected Area', 'Proponent', 'Enumerator', 'Remarks'])
        for report in reports:
            ws.append([report.get('id'), str(report.get('report_date') or ''), report.get('establishment_name') or '', report.get('establishment_type') or '', report.get('establishment_status') or '', report.get('pa_name') or '', report.get('proponent_name') or '', report.get('enumerator_name') or '', report.get('remarks') or ''])
        buffer = BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        response = HttpResponse(buffer.read(), content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = 'attachment; filename="reports.xlsx"'
        return response
    
    elif format_type == 'word':
        from docx import Document
        doc = Document()
        doc.add_heading('Enumerator Reports', 0)
        table = doc.add_table(rows=1, cols=9)
        table.style = 'Light Grid Accent 1'
        hdr = table.rows[0].cells
        headers = ['Report ID', 'Date', 'Establishment', 'Type', 'Status', 'Protected Area', 'Proponent', 'Enumerator', 'Remarks']
        for i, h in enumerate(headers):
            hdr[i].text = h
        for report in reports:
            row = table.add_row().cells
            row[0].text = str(report.get('id') or '')
            row[1].text = str(report.get('report_date') or '')
            row[2].text = report.get('establishment_name') or ''
            row[3].text = report.get('establishment_type') or ''
            row[4].text = report.get('establishment_status') or ''
            row[5].text = report.get('pa_name') or ''
            row[6].text = report.get('proponent_name') or ''
            row[7].text = report.get('enumerator_name') or ''
            row[8].text = report.get('remarks') or ''
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        response = HttpResponse(buffer.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename="reports.docx"'
        return response
    
    else:  # PDF
        from reportlab.lib.pagesizes import letter, landscape
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
        from reportlab.lib.styles import getSampleStyleSheet
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=landscape(letter))
        elements = []
        data = [['ID', 'Date', 'Establishment', 'Type', 'Status', 'PA', 'Proponent', 'Enumerator', 'Remarks']]
        for report in reports:
            data.append([str(report.get('id', '')), str(report.get('report_date', '')), (report.get('establishment_name') or '')[:20], (report.get('establishment_type') or '')[:15], (report.get('establishment_status') or '')[:15], (report.get('pa_name') or '')[:15], (report.get('proponent_name') or '')[:20], (report.get('enumerator_name') or '')[:20], (report.get('remarks') or '')[:30]])
        table = Table(data)
        table.setStyle(TableStyle([('BACKGROUND', (0, 0), (-1, 0), colors.grey), ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke), ('ALIGN', (0, 0), (-1, -1), 'CENTER'), ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'), ('FONTSIZE', (0, 0), (-1, 0), 10), ('BOTTOMPADDING', (0, 0), (-1, 0), 12), ('BACKGROUND', (0, 1), (-1, -1), colors.beige), ('GRID', (0, 0), (-1, -1), 1, colors.black)]))
        elements.append(table)
        doc.build(elements)
        buffer.seek(0)
        response = HttpResponse(buffer.read(), content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename="reports.pdf"'
        return response


# =========================
# USER MANAGEMENT FUNCTIONS
# =========================
def get_all_users():
    """Fetch all users from database."""
    users = []
    try:
        with connection.cursor() as cur:
            cur.execute("""
                SELECT id, username, first_name, last_name, email, role, 
                       region_id, penro_id, cenro_id
                FROM users
                ORDER BY created_at DESC;
            """)
            columns = [desc[0] for desc in cur.description]
            for row in cur.fetchall():
                user_dict = dict(zip(columns, row))
                user_dict['user_id'] = user_dict['id']
                user_dict['office'] = 'N/A'
                if user_dict.get('cenro_id'):
                    user_dict['office'] = 'CENRO'
                elif user_dict.get('penro_id'):
                    user_dict['office'] = 'PENRO'
                elif user_dict.get('region_id'):
                    user_dict['office'] = 'Region'
                users.append(user_dict)
    except DatabaseError as e:
        logger.error(f"Error fetching users: {e}")
    return users


def update_user_profile(user_id, first_name, last_name, username, email=None, role=None):
    """Update user profile information."""
    try:
        with connection.cursor() as cur:
            if email and role:
                # Admin updating user
                cur.execute("""
                    UPDATE users
                    SET first_name = %s, last_name = %s, username = %s, 
                        email = %s, role = %s
                    WHERE id = %s;
                """, [first_name, last_name, username, email, role, user_id])
            else:
                # User updating own profile
                cur.execute("""
                    UPDATE users
                    SET first_name = %s, last_name = %s, username = %s
                    WHERE id = %s;
                """, [first_name, last_name, username, user_id])
        return True, "Profile updated successfully"
    except DatabaseError as e:
        logger.error(f"Error updating user profile: {e}")
        return False, str(e)


def change_user_password(user_id, current_password, new_password):
    """Change user password after verifying current password."""
    try:
        with connection.cursor() as cur:
            # Verify current password
            cur.execute("""
                SELECT password FROM users WHERE id = %s;
            """, [user_id])
            row = cur.fetchone()
            if not row:
                return False, "User not found"
            
            stored_password = row[0]
            # Simple comparison - in production, use proper password hashing
            if stored_password != current_password:
                return False, "Current password is incorrect"
            
            # Update password
            cur.execute("""
                UPDATE users SET password = %s WHERE id = %s;
            """, [new_password, user_id])
        
        return True, "Password changed successfully"
    except DatabaseError as e:
        logger.error(f"Error changing password: {e}")
        return False, str(e)


def delete_user(user_id):
    """Delete user from database."""
    try:
        with connection.cursor() as cur:
            cur.execute("DELETE FROM users WHERE id = %s;", [user_id])
        return True, "User deleted successfully"
    except DatabaseError as e:
        logger.error(f"Error deleting user: {e}")
        return False, str(e)


def get_user_phone(user_id):
    """Get user's phone number."""
    try:
        with connection.cursor() as cur:
            cur.execute("SELECT phone_number FROM users WHERE id = %s;", [user_id])
            row = cur.fetchone()
            return row[0] if row else ''
    except DatabaseError as e:
        logger.error(f"Error fetching phone number: {e}")
        return ''


def update_user_profile_with_phone(user_id, first_name, last_name, username, phone_number=None):
    """Update user profile including phone number."""
    try:
        with connection.cursor() as cur:
            cur.execute("""
                UPDATE users
                SET first_name = %s, last_name = %s, username = %s, phone_number = %s
                WHERE id = %s;
            """, [first_name, last_name, username, phone_number, user_id])
        return True, "Profile updated successfully"
    except DatabaseError as e:
        logger.error(f"Error updating user profile: {e}")
        return False, str(e)


def update_user_profile_with_phone_and_pic(user_id, first_name, last_name, phone_number=None, profile_pic=None):
    """Update user profile including phone number and profile picture."""
    try:
        profile_pic_url = None
        if profile_pic:
            # Upload to Supabase
            bucket = os.getenv('SUPABASE_BUCKET', 'geo-tagged-photos')
            filename = f"profile_pics/{user_id}_{int(time.time())}_{profile_pic.name}"
            file_data = profile_pic.read()
            supabase.storage.from_(bucket).upload(filename, file_data, {"content-type": profile_pic.content_type})
            profile_pic_url = filename
        
        with connection.cursor() as cur:
            if profile_pic_url:
                cur.execute("""
                    UPDATE users
                    SET first_name = %s, last_name = %s, phone_number = %s, profile_pic = %s
                    WHERE id = %s;
                """, [first_name, last_name, phone_number, profile_pic_url, user_id])
            else:
                cur.execute("""
                    UPDATE users
                    SET first_name = %s, last_name = %s, phone_number = %s
                    WHERE id = %s;
                """, [first_name, last_name, phone_number, user_id])
        return True, "Profile updated successfully"
    except Exception as e:
        logger.error(f"Error updating user profile: {e}")
        return False, str(e)


def get_user_profile_pic(user_id):
    """Get user's profile picture URL."""
    try:
        with connection.cursor() as cur:
            cur.execute("SELECT profile_pic FROM users WHERE id = %s;", [user_id])
            row = cur.fetchone()
            if row and row[0]:
                bucket = os.getenv('SUPABASE_BUCKET', 'geo-tagged-photos')
                return f"{SUPABASE_URL}/storage/v1/object/public/{bucket}/{row[0]}"
            return None
    except DatabaseError as e:
        logger.error(f"Error fetching profile picture: {e}")
        return None
